"""Canonical Phase 4 end-to-end run for R-1001.

Runs the LangGraph maintenance agent under the network guard, then evaluates
the result against the (agent-inaccessible) ground truth and writes the report.
"""
import json
import logging
import os
import sys

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")

HERE = os.path.dirname(os.path.abspath(__file__))
BACKEND = os.path.dirname(HERE)
if BACKEND not in sys.path:
    sys.path.insert(0, BACKEND)

from agent.run import run_agent_task
from agent.evaluation.evaluate import evaluate, write_report
from docx import Document

TASK = (
    "Analyze the recent R-1001 operating data and inspection findings, compare them "
    "with the equipment manual and maintenance SOP, check the vendor recommendation, "
    "determine the required corrective action, and prepare a maintenance approval note."
)

# New version so the previously-generated artifact is NOT overwritten.
NEW_ARTIFACT = "R-1001_agent_maintenance_approval_v2.docx"

# Elements that must appear in the new DOCX to prove the complete evidence chain.
EXPECTED_FRAGMENTS = [
    "Sensor Evidence Table",
    "reactor temperature", "reactor pressure", "reactor vibration",
    "HIGH", "Breach",
    "Catalyst Hotspot", "Thermowell Reading Drift", "Top-Head Gasket Weep",
    "HRS-CAT-22", "HRS-GSK-1001", "HRS-TW-1001",
    "Controlled", "Corrective", "APPROVAL REQUIRED",
    "Synthetic", "source",
]


def _docx_full_text(path: str) -> str:
    """Extract ALL text including table cells (verify_docx only reads paragraphs)."""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def verify_complete_evidence(path: str):
    text = _docx_full_text(path)
    lowered = text.lower()
    missing = [frag for frag in EXPECTED_FRAGMENTS if frag.lower() not in lowered]
    return {"path": path, "ok": not missing, "missing": missing, "chars": len(text)}


def main():
    result = run_agent_task(TASK, asset_tag="R-1001", artifact_filename=NEW_ARTIFACT)

    print("\n===== AGENT RUN =====")
    print("run_id:", result["run_id"])
    print("status:", result["status"])
    print("decision:", result["decision"])
    print("approval_required:", result["approval_required"])
    print("artifacts:", result["artifacts"])
    print("external_calls:", result["external_calls"])
    print("errors:", result["errors"])
    print("verification:", result.get("verification"))
    print("evidence_count:", len(result["evidence"]))
    print("trace_nodes:", [t.get("node") for t in result.get("trace", [])])

    artifact = result["artifacts"][0] if result["artifacts"] else None
    evidence_check = verify_complete_evidence(artifact) if artifact else {"ok": False}
    print("\n===== COMPLETE-EVIDENCE CHECK =====")
    print("artifact:", evidence_check.get("path"))
    print("evidence_ok:", evidence_check["ok"])
    print("missing_fragments:", evidence_check.get("missing"))

    report = evaluate(result)
    path = write_report(report, result)
    print("\n===== EVALUATION =====")
    print(f"score: {report['passed']}/{report['total']} ({report['score']}%)")
    print(f"external_calls: {report['external_calls']}")
    print("report:", path)
    print(json.dumps(report["criteria"], indent=2))

    assert result["status"] == "VERIFIED", result["status"]
    assert result["external_calls"] == 0, "network sovereignty violated"
    assert evidence_check["ok"], f"missing evidence: {evidence_check['missing']}"
    assert report["passed"] == report["total"], "evaluation failed"
    print("\nALL CHECKS PASSED")


if __name__ == "__main__":
    main()
