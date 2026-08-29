"""Tool: create_docx

Generates the maintenance approval note DOCX with the complete evidence chain
returned by the agent.

The renderer is fully data-driven: it renders whatever structured evidence the
agent produced (sensor table, inspection findings, vendor parts, SOP requirements,
evidence provenance). It does NOT hard-code any findings.

Writes to the agent output directory (never the original synthetic approval note).
"""
import logging
import os
from datetime import datetime
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agent.config import OUTPUT_DIR, AGENT_AUTHOR

logger = logging.getLogger(__name__)

# The 12 sections required by the Phase 4 specification (in document order).
REQUIRED_SECTIONS = [
    "Title", "Asset Information", "Executive Summary", "Evidence Reviewed",
    "Sensor Findings", "Inspection Findings", "SOP/Manual Requirements",
    "Vendor Recommendation", "Corrective Action", "Approval Request",
    "Source References", "Synthetic Data Disclaimer",
]

DISCLAIMER = (
    "DISCLAIMER: This document was generated autonomously by the Sovereign AI "
    "maintenance agent from SYNTHETIC demonstration data (R-1001 demo dataset). "
    "It is for demonstration and evaluation only and must not be used for real "
    "plant operations, safety decisions, or regulatory submissions."
)


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _kv(doc, key, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{key}: ")
    r.bold = True
    p.add_run(str(value))


def _table(doc, headers: List[str], rows: List[List[str]], widths: Optional[List[float]] = None):
    """Add a styled table. Row cells are rendered as plain text."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def _inspection_label(ftype: str) -> str:
    return {
        "catalyst_hotspot": "Catalyst Hotspot / Deactivation",
        "thermowell_drift": "Thermowell Reading Drift",
        "gasket_weep": "Top-Head Gasket Weep / Seepage",
        "observation": "Observation",
    }.get(ftype, ftype.replace("_", " ").title())


def create_approval_note(content: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """Build the approval note with the complete evidence chain.

    ``output_path`` may be a full path, a bare filename (joined to OUTPUT_DIR), or
    None (defaults to the canonical R-1001 approval note filename).

    Returns the output file path.
    """
    asset_tag = content.get("asset_tag", "R-1001")
    doc = Document()

    # 1. Title
    _h(doc, content.get("title", "R-1001 Maintenance Approval Note"), level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Prepared by {AGENT_AUTHOR}")
    r.italic = True
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    _kv(doc, "Status", "DRAFT — pending human authorization")

    # 2. Asset Information
    _h(doc, "Asset Information", level=1)
    _kv(doc, "Asset Tag", asset_tag)
    for k, v in (content.get("asset_information") or {}).items():
        if k == "Asset Tag":
            continue
        _kv(doc, k, v)

    # 3. Executive Summary
    _h(doc, "Executive Summary", level=1)
    doc.add_paragraph(content.get("executive_summary", ""))

    # 4. Evidence Reviewed (raw retrieved evidence with provenance)
    _h(doc, "Evidence Reviewed", level=1)
    for line in content.get("evidence_reviewed", []):
        doc.add_paragraph(str(line), style="List Bullet")
    chain = content.get("evidence_chain") or []
    if chain:
        _table(
            doc,
            ["Document Type", "Claim / Category", "Confidence", "Source File"],
            [[c.get("document_type", ""), c.get("claim", ""),
              c.get("confidence", ""), c.get("source_file", "")] for c in chain],
            widths=[1.3, 1.6, 0.9, 2.2],
        )

    # 5. Sensor Findings (keyword paragraphs + structured evidence table)
    _h(doc, "Sensor Findings", level=1)
    for b in content.get("sensor_findings", []):
        doc.add_paragraph(str(b), style="List Bullet")
    note = content.get("sensor_note")
    if note:
        doc.add_paragraph(note)
    table_rows = content.get("sensor_table") or []
    if table_rows:
        _h(doc, "Sensor Evidence Table", level=2)
        _table(
            doc,
            ["Signal", "Observed Value", "Threshold", "Timestamp", "Breach Status"],
            [[r.get("signal", ""), r.get("observed_value", ""), r.get("threshold", ""),
              r.get("timestamp", ""), r.get("breach_status", "")] for r in table_rows],
            widths=[1.6, 1.7, 1.5, 1.2, 1.7],
        )

    # 6. Inspection Findings (catalyst hotspot / thermowell drift / gasket weep)
    _h(doc, "Inspection Findings", level=1)
    insp = content.get("inspection_list") or []
    if insp:
        for item in insp:
            p = doc.add_paragraph()
            p.add_run(_inspection_label(item.get("type", "")) + ": ").bold = True
            p.add_run(str(item.get("value", "")))
            src = item.get("source")
            if src:
                sp = doc.add_paragraph()
                sr = sp.add_run(f"Source: {src}")
                sr.italic = True
    else:
        for b in content.get("inspection_findings", []):
            doc.add_paragraph(str(b), style="List Bullet")

    # 7. SOP/Manual Requirements
    _h(doc, "SOP/Manual Requirements", level=1)
    for b in content.get("sop_requirements", []):
        doc.add_paragraph(str(b), style="List Bullet")

    # 8. Vendor Recommendation
    _h(doc, "Vendor Recommendation", level=1)
    vlist = content.get("vendor_list") or []
    if vlist:
        _table(
            doc,
            ["Recommended Spare Part", "Source"],
            [[v.get("part_number", ""), v.get("source", "")] for v in vlist],
            widths=[2.5, 2.5],
        )
    else:
        for b in content.get("vendor_recommendation", []):
            doc.add_paragraph(str(b), style="List Bullet")

    # 9. Corrective Action
    _h(doc, "Corrective Action", level=1)
    for b in content.get("corrective_action", []):
        doc.add_paragraph(str(b), style="List Number")

    # 10. Approval Request
    _h(doc, "Approval Request", level=1)
    doc.add_paragraph(content.get("approval_request", ""))
    p = doc.add_paragraph()
    if content.get("approval_required"):
        p.add_run("APPROVAL REQUIRED: YES").bold = True
    else:
        p.add_run("APPROVAL REQUIRED: NO").bold = True

    # 11. Source References
    _h(doc, "Source References", level=1)
    for s in content.get("source_references", []):
        doc.add_paragraph(str(s), style="List Bullet")

    # 12. Synthetic Data Disclaimer
    _h(doc, "Synthetic Data Disclaimer", level=1)
    doc.add_paragraph(DISCLAIMER)

    # Resolve output path (never overwrite the original synthetic note).
    if output_path:
        if os.path.isabs(output_path):
            out_path = output_path
        else:
            out_path = os.path.join(str(OUTPUT_DIR), output_path)
    else:
        out_path = os.path.join(str(OUTPUT_DIR), f"{asset_tag}_agent_maintenance_approval.docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    logger.info("Created approval note: %s", out_path)
    return out_path


def verify_docx(path: str) -> Dict[str, Any]:
    """Re-open a DOCX and verify all required sections/content are present."""
    missing_sections = []
    asset_present = False
    disclaimer_present = False
    sources_present = False
    try:
        doc = Document(path)
    except Exception as e:
        return {"ok": False, "error": f"cannot open: {e}", "missing_sections": REQUIRED_SECTIONS}

    full_text = "\n".join(p.text for p in doc.paragraphs)
    lowered = full_text.lower()

    section_markers = {
        "Title": None,  # implicit (title heading)
        "Asset Information": "asset information",
        "Executive Summary": "executive summary",
        "Evidence Reviewed": "evidence reviewed",
        "Sensor Findings": "sensor findings",
        "Inspection Findings": "inspection finding",
        "SOP/Manual Requirements": "sop",
        "Vendor Recommendation": "vendor recommendation",
        "Corrective Action": "corrective action",
        "Approval Request": "approval request",
        "Source References": "source reference",
        "Synthetic Data Disclaimer": "disclaimer",
    }
    for label, marker in section_markers.items():
        if label == "Title":
            continue
        if marker and marker not in lowered:
            missing_sections.append(label)

    if "R-1001" in full_text:
        asset_present = True
    else:
        missing_sections.append("Asset Tag (R-1001)")

    if "disclaimer" in lowered and "synthetic" in lowered:
        disclaimer_present = True
    else:
        missing_sections.append("Synthetic Data Disclaimer")

    if "source reference" in lowered or "source:" in lowered or "assets/r-1001" in lowered.replace(" ", ""):
        sources_present = True
    else:
        missing_sections.append("Source References")

    ok = (not missing_sections) and asset_present and disclaimer_present and sources_present
    return {
        "ok": ok,
        "missing_sections": missing_sections,
        "asset_present": asset_present,
        "disclaimer_present": disclaimer_present,
        "sources_present": sources_present,
        "paragraphs": len(doc.paragraphs),
    }
