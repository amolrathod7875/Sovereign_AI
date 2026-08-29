"""Document parsers: DOCX, PDF, XLSX, CSV, JSON, EML -> structured blocks.

Each parser returns a list of blocks:
  {"type": "heading"|"para"|"table"|"row", "text": str, "meta": {...}}
The text is retrieval-friendly; structure (headings/sections/tables) is preserved
in the block type so the chunker can keep it coherent.
"""
import json
import os
import csv
from typing import List, Dict, Any


def _blocks_docx(path: str) -> List[Dict[str, Any]]:
    from docx import Document
    d = Document(path)
    blocks = []
    for p in d.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = (p.style.name or "").lower()
        btype = "heading" if style.startswith("heading") else "para"
        blocks.append({"type": btype, "text": txt})
    # tables as their own blocks
    for ti, tbl in enumerate(d.tables):
        rows = []
        for r in tbl.rows:
            rows.append(" | ".join(c.text.strip() for c in r.cells))
        blocks.append({"type": "table", "text": f"[Table {ti+1}]\n" + "\n".join(rows)})
    return blocks


def _blocks_pdf(path: str) -> List[Dict[str, Any]]:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    blocks = []
    for page in doc:
        d = page.get_text("dict")
        for b in d.get("blocks", []):
            if b["type"] != 0:
                continue
            for line in b.get("lines", []):
                txt = " ".join(s["text"] for s in line["spans"]).strip()
                if not txt:
                    continue
                size = max((s["size"] for s in line["spans"]), default=10)
                btype = "heading" if size >= 13 else "para"
                blocks.append({"type": btype, "text": txt})
    return blocks


def _blocks_xlsx(path: str) -> List[Dict[str, Any]]:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks = []
    for ws in wb.worksheets:
        blocks.append({"type": "heading", "text": f"Sheet: {ws.title}"})
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        for ri, row in enumerate(rows[1:], start=1):
            cells = [str(c) if c is not None else "" for c in row]
            blocks.append({"type": "row",
                           "text": f"Row {ri}: " + " | ".join(f"{h}={v}" for h, v in zip(header, cells)),
                           "meta": {"row": ri}})
    return blocks


def _blocks_csv(path: str) -> List[Dict[str, Any]]:
    blocks = []
    with open(path, encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        header = reader.fieldnames or []
        for ri, row in enumerate(reader, start=1):
            cells = {h: (row.get(h) or "") for h in header}
            blocks.append({"type": "row",
                           "text": f"Row {ri}: " + " | ".join(f"{h}={v}" for h, v in cells.items()),
                           "meta": {"row": ri}})
    return blocks


def _blocks_json(path: str) -> List[Dict[str, Any]]:
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    blocks = []

    def walk(node, prefix=""):
        if isinstance(node, dict):
            for k, v in node.items():
                if k == "data_origin":
                    continue
                if isinstance(v, (dict, list)):
                    blocks.append({"type": "heading", "text": prefix + str(k)})
                    walk(v, prefix + str(k) + " / ")
                else:
                    blocks.append({"type": "para", "text": f"{prefix}{k}: {v}"})
        elif isinstance(node, list):
            for i, v in enumerate(node):
                if isinstance(v, (dict, list)):
                    walk(v, prefix + f"[{i}] ")
                else:
                    blocks.append({"type": "para", "text": f"{prefix}[{i}]: {v}"})
        else:
            blocks.append({"type": "para", "text": f"{prefix}: {node}"})

    walk(data)
    return blocks


def _blocks_eml(path: str) -> List[Dict[str, Any]]:
    with open(path, encoding="utf-8", errors="replace") as f:
        raw = f.read()
    blocks = [{"type": "heading", "text": "Vendor Correspondence (EML)"}]
    # crude header/body split
    if "\n\n" in raw:
        head, body = raw.split("\n\n", 1)
    else:
        head, body = raw, ""
    blocks.append({"type": "para", "text": head.strip()})
    if body.strip():
        blocks.append({"type": "para", "text": body.strip()})
    return blocks


def parse_file(path: str) -> List[Dict[str, Any]]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _blocks_docx(path)
    if ext == ".pdf":
        return _blocks_pdf(path)
    if ext == ".xlsx":
        return _blocks_xlsx(path)
    if ext == ".csv":
        return _blocks_csv(path)
    if ext == ".json":
        return _blocks_json(path)
    if ext == ".eml":
        return _blocks_eml(path)
    raise ValueError(f"Unsupported file type: {ext}")
