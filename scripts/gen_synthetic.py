"""Phase 2 - Synthetic Industrial Knowledge Dataset for R-1001 (Reactor, 158.jpg).

Locally generated, deterministic (seed=42), no cloud.
Source of truth for identity/context: verified_equipment_registry.json (R-1001, Reactor,
Hydrogen Production Plant, 158.jpg). All other values are explicitly synthetic_demo.
"""
import base64 as _b  # noqa (unused, keeps nothing)
import csv
import json
import os
import random
import textwrap
from datetime import datetime, timedelta

ROOT = r"D:\Sovereign_AI\data\synthetic"
ASSET = os.path.join(ROOT, "assets", "R-1001")
for p in [
    os.path.join(ROOT, "plant"),
    os.path.join(ROOT, "metadata"),
    os.path.join(ASSET, "manual"),
    os.path.join(ASSET, "sop"),
    os.path.join(ASSET, "inspection"),
    os.path.join(ASSET, "maintenance"),
    os.path.join(ASSET, "sensors"),
    os.path.join(ASSET, "correspondence"),
    os.path.join(ASSET, "approvals"),
]:
    os.makedirs(p, exist_ok=True)

random.seed(42)
GEN_TS = datetime(2026, 8, 29, 7, 16, 0)  # fixed for reproducibility of narrative dates
GEN_ISO = GEN_TS.isoformat(timespec="seconds")

# ---------------------------------------------------------------- profile
PROFILE = {
    "profile_version": "1.0.0",
    "generated_utc": GEN_ISO,
    "public_pid_identity": {
        "data_origin": "public_pid",
        "asset_tag": "R-1001",
        "equipment_type": "Reactor",
        "source_drawing": "158.jpg",
        "plant": "Hydrogen Production Plant",
        "verification_status": "verified",
        "verified_source": "data/pid_analysis/registry/verified_equipment_registry.json",
        "note": "Identity/context taken ONLY from the verified P&ID. No extra real-plant facts invented.",
    },
    "synthetic_demo_profile": {
        "data_origin": "synthetic_demo",
        "service": "Catalytic hydrogen-synthesis reactor (DEMO service description; not a real plant fact)",
        "manufacturer": "HydroReactor Systems GmbH (SYNTHETIC vendor, demo only)",
        "model": "HRS-R1000 (SYNTHETIC model no., demo only)",
        "tag_on_pid": "R-1001",
        "design": {"pressure_bar": 25.0, "temperature_c": 350.0, "data_origin": "synthetic_demo"},
        "operating_parameters": {
            "pressure_bar": 18.0, "temperature_c": 280.0, "level_pct": 60.0,
            "throughput_nm3_h": 4200.0, "data_origin": "synthetic_demo",
        },
        "alarm_thresholds": {
            "temp_high_c": 310.0, "temp_high_high_c": 320.0, "temp_low_c": 250.0,
            "pressure_high_bar": 21.0, "pressure_high_high_bar": 25.0,
            "vibration_high_mm_s": 4.0, "vibration_high_high_mm_s": 6.0,
            "level_low_pct": 20.0, "level_high_pct": 90.0,
            "data_origin": "synthetic_demo",
        },
        "maintenance_intervals": {
            "reactor_full_inspection_months": 12, "catalyst_activity_test_months": 6,
            "thermowell_calibration_months": 6, "gasket_visual_check_months": 12,
            "data_origin": "synthetic_demo",
        },
        "inspection_criteria": {
            "max_vibration_mm_s": 4.0, "max_shell_temp_gradient_c": 15.0,
            "gasket_condition_ok": "no weep / no crust", "catalyst_bed_temp_uniformity_c": 10.0,
            "thermowell_drift_max_c": 2.0, "data_origin": "synthetic_demo",
        },
        "failure_modes": [
            {"mode": "Catalyst deactivation leading to thermal hotspot", "likelihood": "medium",
             "data_origin": "synthetic_demo"},
            {"mode": "Thermowell reading drift / sensor fault", "likelihood": "medium",
             "data_origin": "synthetic_demo"},
            {"mode": "Top-head gasket weep / hydrogen leak", "likelihood": "low",
             "data_origin": "synthetic_demo"},
            {"mode": "Internal corrosion / wall thinning", "likelihood": "low",
             "data_origin": "synthetic_demo"},
        ],
        "spare_parts": [
            {"part": "Catalyst charge (Ni-based)", "pn": "HRS-CAT-22", "data_origin": "synthetic_demo"},
            {"part": "Top-head gasket kit", "pn": "HRS-GSK-1001", "data_origin": "synthetic_demo"},
            {"part": "Reactor thermowell assembly", "pn": "HRS-TW-1001", "data_origin": "synthetic_demo"},
            {"part": "Cooling coil bundle section", "pn": "HRS-CC-1001", "data_origin": "synthetic_demo"},
        ],
        "sensor_definitions": [
            {"tag": "TI-1001", "measurement": "Reactor temperature", "unit": "C", "range": [0, 400],
             "alarm_high": 310.0, "alarm_high_high": 320.0, "data_origin": "synthetic_demo"},
            {"tag": "PI-1001", "measurement": "Reactor pressure", "unit": "bar", "range": [0, 30],
             "alarm_high": 21.0, "alarm_high_high": 25.0, "data_origin": "synthetic_demo"},
            {"tag": "VI-1001", "measurement": "Reactor vibration (bearing housing)", "unit": "mm/s",
             "range": [0, 10], "alarm_high": 4.0, "alarm_high_high": 6.0,
             "data_origin": "synthetic_demo"},
            {"tag": "LI-1001", "measurement": "Reactor level", "unit": "%", "range": [0, 100],
             "alarm_low": 20.0, "alarm_high": 90.0, "data_origin": "synthetic_demo"},
        ],
        "maintenance_history_summary": [
            {"wo": "WO-2023-014", "date": "2023-06-12", "type": "Preventive",
             "summary": "12-month reactor inspection; no defect", "data_origin": "synthetic_demo"},
            {"wo": "WO-2024-021", "date": "2024-06-09", "type": "Preventive",
             "summary": "12-month inspection + catalyst activity test (79% activity)", "data_origin": "synthetic_demo"},
            {"wo": "WO-2025-019", "date": "2025-06-15", "type": "Preventive",
             "summary": "12-month reactor inspection; gasket OK", "data_origin": "synthetic_demo"},
        ],
    },
}
with open(os.path.join(ASSET, "profile.json"), "w", encoding="utf-8") as f:
    json.dump(PROFILE, f, indent=2, ensure_ascii=False)

SP = PROFILE["synthetic_demo_profile"]
print("[ok] profile.json written")

# ---------------------------------------------------------------- sensor dataset (deterministic)
SENSOR_CSV = os.path.join(ASSET, "sensors", "sensor_dataset.csv")
DAYS = 90
HOURS = DAYS * 24
END = GEN_TS
START = END - timedelta(hours=HOURS - 1)
DRIFT_START_H = 55 * 24  # degradation begins at hour 1320 (~day 55)

def gen_series():
    rows = []
    ti_h, ti_hh = SP["alarm_thresholds"]["temp_high_c"], SP["alarm_thresholds"]["temp_high_high_c"]
    pi_h = SP["alarm_thresholds"]["pressure_high_bar"]
    vi_h = SP["alarm_thresholds"]["vibration_high_mm_s"]
    for h in range(HOURS):
        ts = START + timedelta(hours=h)
        if h < DRIFT_START_H:
            temp = 280.0 + random.gauss(0, 0.4)
            press = 18.0 + random.gauss(0, 0.1)
            vib = 2.1 + random.gauss(0, 0.15)
            lvl = 60.0 + random.gauss(0, 1.5)
        else:
            dh = h - DRIFT_START_H
            temp = 280.0 + 0.05 * dh + random.gauss(0, 0.5)
            press = 18.0 + 0.005 * dh + random.gauss(0, 0.12)
            vib = 2.1 + 0.0035 * dh + random.gauss(0, 0.18)
            lvl = 60.0 + random.gauss(0, 1.5)
        temp = round(temp, 2); press = round(press, 3); vib = round(vib, 3); lvl = round(lvl, 1)
        anomaly = (temp >= ti_h) or (press >= pi_h) or (vib >= vi_h)
        rows.append({
            "timestamp": ts.strftime("%Y-%m-%dT%H:%M:%S"),
            "TI-1001_reactor_temp_C": temp,
            "PI-1001_reactor_pressure_bar": press,
            "VI-1001_reactor_vibration_mm_s": vib,
            "LI-1001_reactor_level_pct": lvl,
            "anomaly_flag": int(anomaly),
        })
    return rows

rows = gen_series()
with open(SENSOR_CSV, "w", encoding="utf-8", newline="") as f:
    w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
    w.writeheader(); w.writerows(rows)
print(f"[ok] sensor_dataset.csv written ({len(rows)} rows)")

# ---------------------------------------------------------------- maintenance history XLSX
import openpyxl
MH_XLSX = os.path.join(ASSET, "maintenance", "maintenance_history.xlsx")
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "WorkOrders"
cols = ["work_order", "date", "type", "description", "part_pn", "technician", "status", "asset", "data_origin"]
ws.append(cols)
mh_rows = [
    ["WO-2023-014", "2023-06-12", "Preventive", "12-month reactor inspection; no defect found",
     "", "J. Okafor", "Closed", "R-1001", "synthetic_demo"],
    ["WO-2024-021", "2024-06-09", "Preventive", "12-month inspection + catalyst activity test (result 79% activity)",
     "HRS-CAT-22", "A. Meena", "Closed", "R-1001", "synthetic_demo"],
    ["WO-2025-019", "2025-06-15", "Preventive", "12-month reactor inspection; top-head gasket OK",
     "", "L. Brandt", "Closed", "R-1001", "synthetic_demo"],
    ["WO-2026-031", "2026-08-27", "Corrective", "Investigative internal inspection after TI-1001 high-high "
     "alarm: catalyst-bed thermal hotspot + thermowell drift + top-head gasket weep identified; "
     "catalyst replacement + gasket replacement + thermowell recalibration recommended",
     "HRS-CAT-22; HRS-GSK-1001; HRS-TW-1001", "L. Brandt", "Pending Approval", "R-1001", "synthetic_demo"],
]
for r in mh_rows:
    ws.append(r)
wb.save(MH_XLSX)
print("[ok] maintenance_history.xlsx written")

# ---------------------------------------------------------------- DOCX helpers
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def new_doc(title, subtitle=None):
    d = Document()
    h = d.add_heading(title, level=0)
    if subtitle:
        p = d.add_paragraph(subtitle)
        p.runs[0].italic = True
    d.add_paragraph("ASSET: R-1001  |  PLANT: Hydrogen Production Plant (158.jpg)  |  "
                    "DATA ORIGIN: synthetic_demo").runs[0].font.size = Pt(9)
    return d

def para(d, text, bold=False, size=None):
    p = d.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if size:
        p.runs[0].font.size = Pt(size)
    return p

def bullets(d, items):
    for it in items:
        d.add_paragraph(it, style="List Bullet")

def save_doc(d, path):
    d.save(path)
    print(f"[ok] {os.path.basename(path)} written ({path})")

# 1. Equipment Manual
man = new_doc("R-1001 Reactor - Equipment Manual", "Synthetic demo manual (not real plant data)")
para(man, "1. Identity", bold=True)
para(man, "Asset tag: R-1001 (from verified P&ID, 158.jpg, Hydrogen Production Plant).")
para(man, "Equipment type: Reactor (verified).")
para(man, "Synthetic manufacturer / model: HydroReactor Systems GmbH, model HRS-R1000 (SYNTHETIC).")
para(man, "Synthetic service: Catalytic hydrogen-synthesis reactor (DEMO description).")
para(man, "2. Design & Operating Limits (synthetic_demo)", bold=True)
bullets(man, [
    "Design pressure: 25.0 bar; Design temperature: 350.0 C.",
    "Normal operating pressure: 18.0 bar; Normal operating temperature: 280.0 C.",
    "Maximum continuous temperature: 320.0 C (high-high alarm); Maximum continuous pressure: 21.0 bar.",
    "Normal level: 60% (alarm low 20%, alarm high 90%).",
    "Throughput (demo): 4200 Nm3/h.",
])
para(man, "3. Alarm Thresholds (synthetic_demo)", bold=True)
bullets(man, [
    "TI-1001 temperature: High 310 C, High-High 320 C, Low 250 C.",
    "PI-1001 pressure: High 21 bar, High-High 25 bar.",
    "VI-1001 vibration: High 4.0 mm/s, High-High 6.0 mm/s.",
])
para(man, "4. Failure Modes & Spares (synthetic_demo)", bold=True)
bullets(man, [
    "Catalyst deactivation -> thermal hotspot; spare catalyst HRS-CAT-22.",
    "Thermowell drift -> spare thermowell HRS-TW-1001.",
    "Top-head gasket weep -> spare gasket kit HRS-GSK-1001.",
    "Internal corrosion -> cooling coil section HRS-CC-1001.",
])
para(man, "5. Maintenance Interval Basis (synthetic_demo)", bold=True)
bullets(man, [
    "Full reactor inspection every 12 months.",
    "Catalyst activity test every 6 months; Thermowell calibration every 6 months.",
])
para(man, "DISCLAIMER: All numeric values in this manual are synthetic demo data. The only "
           "public-P&ID-derived facts are the asset tag R-1001, its type (Reactor), the "
           "drawing 158.jpg, and the plant name Hydrogen Production Plant.", size=9)
save_doc(man, os.path.join(ASSET, "manual", "manual.docx"))

# 2. Operating SOP
op = new_doc("R-1001 Reactor - Operating SOP", "Synthetic demo operating procedure")
para(op, "Purpose: safe startup, operation and normal shutdown of R-1001 within synthetic limits.")
para(op, "1. Normal Operation (synthetic_demo limits)", bold=True)
bullets(op, ["Maintain TI-1001 between 250 C and 310 C (target 280 C).",
             "Maintain PI-1001 at/below 21 bar (target 18 bar).",
             "Maintain VI-1001 below 4.0 mm/s; LI-1001 between 20% and 90%."])
para(op, "2. Abnormal Operation - Temperature", bold=True)
bullets(op, [
    "TI-1001 >= 310 C (HIGH): reduce feed rate 10%, increase cooling; notify shift engineer; "
    "review sensor_dataset.csv trend.",
    "TI-1001 >= 320 C (HIGH-HIGH): initiate automatic reactor trip / ESD; begin emergency "
    "shutdown checklist; do NOT reset until cause identified.",
    "TI-1001 <= 250 C (LOW): risk of side reactions; reduce cooling, verify catalyst activity.",
])
para(op, "3. Abnormal Operation - Vibration / Pressure", bold=True)
bullets(op, [
    "VI-1001 >= 4.0 mm/s: schedule bearing/alignment inspection; restrict load.",
    "PI-1001 >= 21 bar: check downstream restriction; prepare controlled depressurisation.",
])
para(op, "4. Cross-reference", bold=True)
para(op, "Breach events must be correlated with assets/R-1001/sensors/sensor_dataset.csv and, if "
         "a defect is suspected, with the Inspection Report and Maintenance SOP before any "
         "corrective work.")
para(op, "DISCLAIMER: synthetic demo SOP.", size=9)
save_doc(op, os.path.join(ASSET, "sop", "operating_sop.docx"))

# 3. Preventive Maintenance SOP
pm = new_doc("R-1001 Reactor - Preventive Maintenance SOP", "Synthetic demo PM procedure")
para(pm, "Intervals (synthetic_demo): full inspection 12 months; catalyst test 6 months; "
          "thermowell calibration 6 months; gasket visual 12 months.")
para(pm, "1. Routine PM Tasks", bold=True)
bullets(pm, [
    "Visual: gasket no weep / no crust; shell temp gradient <= 15 C.",
    "Verify TI-1001 vs portable reference; thermowell drift <= 2.0 C.",
    "Vibration survey: VI-1001 <= 4.0 mm/s.",
    "Catalyst activity test >= 70% (else plan replacement, P/N HRS-CAT-22).",
])
para(pm, "2. Corrective Action Triggered by Inspection/Alarms (synthetic_demo)", bold=True)
bullets(pm, [
    "Catalyst hotspot / activity < 70%: plan catalyst change-out (HRS-CAT-22) during shutdown.",
    "Thermowell drift > 2.0 C: recalibrate / replace (HRS-TW-1001).",
    "Gasket weep: replace top-head gasket (HRS-GSK-1001); requires controlled shutdown.",
    "Any High-High temperature breach: ESD + shutdown per Operating SOP before corrective work.",
])
para(pm, "3. Approval Gate", bold=True)
para(pm, "Corrective work requiring reactor shutdown and/or spare parts > threshold must be "
         "captured in a Maintenance Approval Note (assets/R-1001/approvals) before execution.")
para(pm, "DISCLAIMER: synthetic demo PM SOP.", size=9)
save_doc(pm, os.path.join(ASSET, "sop", "pm_sop.docx"))

# 5. Maintenance History (DOCX summary mirror, optional) - skip; XLSX is the canonical one
# We already wrote XLSX. Add a short DOCX summary for traceability.
mhd = new_doc("R-1001 Reactor - Maintenance History (summary)", "Synthetic demo history; full log in maintenance_history.xlsx")
bullets(mhd, [
    "WO-2023-014 (2023-06-12) Preventive - 12-month inspection, no defect.",
    "WO-2024-021 (2024-06-09) Preventive - inspection + catalyst test 79% activity.",
    "WO-2025-019 (2025-06-15) Preventive - inspection, gasket OK.",
    "WO-2026-031 (2026-08-27) Corrective - PENDING APPROVAL: hotspot + thermowell drift + "
    "gasket weep found after TI-1001 HH alarm; catalyst+gasket+thermowell work recommended.",
])
para(mhd, "DISCLAIMER: synthetic demo history.", size=9)
save_doc(mhd, os.path.join(ASSET, "maintenance", "maintenance_history.docx"))

# 7. Vendor Correspondence (EML)
EML = os.path.join(ASSET, "correspondence", "vendor_correspondence.eml")
eml_body = f"""Subject: RE: R-1001 temperature & vibration trend - recommendation
From: HydroReactor Systems GmbH <support@hydroreactor.example> (SYNTHETIC vendor)
To: Hydrogen Plant Maintenance Lead <maintenance@plant.example>
Date: {GEN_TS.strftime('%a, %d %b %Y %H:%M:%S +0000')}
Message-ID: <vend-R1001-{GEN_TS.strftime('%Y%m%d')}@hydroreactor.example>
X-Data-Origin: synthetic_demo

Dear Maintenance Lead,

Thank you for sharing the R-1001 sensor trend (TI-1001, PI-1001, VI-1001) covering the
last 90 days. The gradual, correlated rise in reactor temperature (approaching the 320 C
high-high), pressure (toward 21 bar) and vibration (above 4.0 mm/s) is consistent with
catalyst deactivation producing a local thermal hotspot, compounded by thermowell reading
drift and a possible top-head gasket weep.

Our recommendation, consistent with your inspection findings:

1. Catalyst: perform a catalyst activity test. If below 70%, replace the charge with
   HRS-CAT-22 (Ni-based). Lead time 4 weeks.
2. Gasket: replace the top-head gasket kit HRS-GSK-1001 during a controlled shutdown.
3. Thermowell: recalibrate, or replace assembly HRS-TW-1001 if drift exceeds 2.0 C.
4. Do not reset/restart the reactor until the High-High temperature event is closed out
   per your Operating SOP (ESD + shutdown).

We can support a shutdown-window planning call on request.

Best regards,
HydroReactor Systems GmbH - Service (SYNTHETIC demo correspondence; not a real vendor)
"""
with open(EML, "w", encoding="utf-8") as f:
    f.write(eml_body)
print("[ok] vendor_correspondence.eml written")

# 8. Approval Note (DOCX)
ap = new_doc("R-1001 Reactor - Maintenance Approval Note", "Synthetic demo approval request")
para(ap, "Request ID: APP-2026-0042   |   Date: 2026-08-29   |   Status: PENDING APPROVAL", bold=True)
para(ap, "Asset: R-1001 (Reactor, Hydrogen Production Plant, 158.jpg).")
para(ap, "1. Evidence Summary", bold=True)
bullets(ap, [
    "Sensor trend (assets/R-1001/sensors/sensor_dataset.csv): TI-1001 approached 320 C high-high "
    "and crossed 310 C high by ~day 80; VI-1001 exceeded 4.0 mm/s by ~day 78; PI-1001 exceeded "
    "21 bar by ~day 80 - correlated multi-signal degradation.",
    "Inspection Report (assets/R-1001/inspection/inspection_report.pdf): catalyst-bed thermal "
    "hotspot, thermowell drift, top-head gasket weep confirmed.",
    "Equipment Manual / PM SOP define limits and corrective actions; Vendor correspondence "
    "(HydroReactor Systems) recommends catalyst (HRS-CAT-22) + gasket (HRS-GSK-1001) replacement "
    "and thermowell recalibration.",
])
para(ap, "2. Proposed Work (synthetic_demo)", bold=True)
bullets(ap, [
    "Controlled reactor shutdown (ESD per Operating SOP) - 5 days.",
    "Replace catalyst charge HRS-CAT-22; replace top-head gasket HRS-GSK-1001; recalibrate/"
    "replace thermowell HRS-TW-1001.",
    "Post-work catalyst activity test >= 70% before restart.",
])
para(ap, "3. Approval Required", bold=True)
para(ap, "This corrective maintenance requires operations approval because it takes R-1001 "
         "offline (production impact) and consumes charged spares. Expected approval: YES, "
         "subject to shutdown-window scheduling.")
para(ap, "DISCLAIMER: synthetic demo approval note; all values synthetic except R-1001 / "
           "Reactor / 158.jpg / Hydrogen Production Plant identity.", size=9)
save_doc(ap, os.path.join(ASSET, "approvals", "approval_note.docx"))

# ---------------------------------------------------------------- inspection PDF (dependency-free)
def _esc(s):
    return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

def write_pdf(path, title, blocks):
    """blocks: list of (kind, text); kind in {'h1','h2','body'. 'bullet'}."""
    pages = []
    cur = []
    y = 800
    def flush():
        nonlocal cur, y
        pages.append(cur); cur = []; y = 800
    for kind, text in blocks:
        size = {"h1": 16, "h2": 12, "body": 10, "bullet": 10}[kind]
        wrap = 95 if size <= 10 else 60
        prefix = "- " if kind == "bullet" else ""
        lines = []
        for para_line in text.split("\n"):
            wrapped = textwrap.wrap(prefix + para_line, wrap) or [""]
            lines.extend(wrapped)
        for ln in lines:
            if y < 60:
                flush()
            cur.append((y, size, ln))
            y -= size + 6
    if cur:
        flush()
    # build PDF
    objs = []
    # 1 catalog, 2 pages parent, then per page: content + page
    n_pages = len(pages)
    # object layout: 1 catalog, 2 pages, for each page: page obj + content obj
    page_obj_ids = []
    content_obj_ids = []
    next_id = 3
    for _ in range(n_pages):
        page_obj_ids.append(next_id); next_id += 1
        content_obj_ids.append(next_id); next_id += 1
    # catalog
    objs.append((1, "<< /Type /Catalog /Pages 2 0 R >>"))
    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    objs.append((2, f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>"))
    for i, pg in enumerate(pages):
        pid = page_obj_ids[i]; cid = content_obj_ids[i]
        stream = ["BT"]
        last_size = 10
        for (yy, sz, ln) in pg:
            if sz != last_size:
                stream.append(f"/F1 {sz} Tf")
                last_size = sz
            stream.append(f"1 0 0 1 50 {yy} Tm")
            stream.append(f"({_esc(ln)}) Tj")
        stream.append("ET")
        content = "\n".join(stream)
        enc = content.encode("latin-1", "replace")
        objs.append((cid, f"<< /Length {len(enc)} >>\nstream\n{content}\nendstream"))
        objs.append((pid, f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                          f"/Resources << /Font << /F1 4 0 R >> >> /Contents {cid} 0 R >>"))
    # font
    fid = next_id
    objs.append((fid, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))
    objs.sort(key=lambda o: o[0])
    # patch font ref
    new_objs = []
    for oid, val in objs:
        if oid == 4:
            new_objs.append((oid, val))
        else:
            new_objs.append((oid, val))
    # re-number: ensure font is last; reference 4 in pages already
    max_id = max(o[0] for o in objs)
    # write
    out = ["%PDF-1.4"]
    offsets = {}
    buf = "%PDF-1.4\n"
    body = ""
    # Build sequentially with offsets
    parts = []
    offset = len(buf.encode("latin-1"))
    # we need offsets per object id
    obj_map = {oid: val for oid, val in objs}
    ids = sorted(obj_map.keys())
    pdf = "%PDF-1.4\n"
    offs = {}
    for oid in ids:
        offs[oid] = len(pdf.encode("latin-1"))
        val = obj_map[oid]
        if val.startswith("<<") and "stream" in val:
            pdf += f"{oid} 0 obj\n{val}\nendobj\n"
        else:
            pdf += f"{oid} 0 obj\n{val}\nendobj\n"
    xref_pos = len(pdf.encode("latin-1"))
    pdf += f"xref\n0 {max_id+1}\n"
    pdf += "0000000000 65535 f \n"
    for oid in range(1, max_id + 1):
        pdf += f"{offs.get(oid,0):010d} 00000 n \n"
    pdf += f"trailer\n<< /Size {max_id+1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF"
    with open(path, "wb") as f:
        f.write(pdf.encode("latin-1", "replace"))
    print(f"[ok] {os.path.basename(path)} written ({path})")

blocks = [
    ("h1", "R-1001 Reactor - Inspection Report"),
    ("body", "Plant: Hydrogen Production Plant (158.jpg)  |  Asset: R-1001 (Reactor, verified)"),
    ("body", "Inspection date: 2026-08-27  |  Inspector: L. Brandt  |  Data origin: synthetic_demo"),
    ("h2", "1. Trigger"),
    ("body", "Routine review of assets/R-1001/sensors/sensor_dataset.csv showed TI-1001 reactor "
             "temperature approaching the 320 C high-high alarm, VI-1001 vibration above 4.0 mm/s, "
             "and PI-1001 pressure above 21 bar - a correlated multi-signal degradation trend."),
    ("h2", "2. Findings"),
    ("bullet", "Catalyst bed thermal hotspot: local bed temperature non-uniformity > 10 C; "
               "activity test later returned 61% (below 70% limit)."),
    ("bullet", "Thermowell drift: TI-1001 reads ~3.5 C above portable reference (limit 2.0 C)."),
    ("bullet", "Top-head gasket weep: dry-crust trace at gasket indicating minor hydrogen weep."),
    ("bullet", "Vibration 4.0-5.0 mm/s consistent with unbalanced/loaded catalyst bed, not bearing fault."),
    ("h2", "3. Assessment"),
    ("body", "Root cause: catalyst deactivation producing a thermal hotspot, with secondary "
             "thermowell drift and gasket weep. Consistent with synthetic failure-mode list in "
             "profile.json and with vendor correspondence recommendation."),
    ("h2", "4. Recommendation"),
    ("bullet", "Replace catalyst charge HRS-CAT-22 during controlled shutdown."),
    ("bullet", "Replace top-head gasket HRS-GSK-1001."),
    ("bullet", "Recalibrate / replace thermowell HRS-TW-1001."),
    ("bullet", "Do not restart until High-High temperature event closed per Operating SOP; "
               "capture in Maintenance Approval Note APP-2026-0042."),
    ("h2", "5. Disclaimer"),
    ("body", "Synthetic demo inspection report. Only R-1001 / Reactor / 158.jpg / Hydrogen "
             "Production Plant are public-P&ID-derived; all findings and values are synthetic."),
]
write_pdf(os.path.join(ASSET, "inspection", "inspection_report.pdf"), "R-1001 Inspection Report", blocks)

# ---------------------------------------------------------------- plant context
PLANT = os.path.join(ROOT, "plant", "plant_context.json")
plant_ctx = {
    "data_origin": "synthetic_demo",
    "plant": "Hydrogen Production Plant",
    "source_drawing": "158.jpg",
    "note": "Only the plant name and drawing 158.jpg are public-P&ID-derived. All other context "
            "below is synthetic demo data used to frame the R-1001 asset chain.",
    "synthetic_context": {
        "process": "Catalytic hydrogen synthesis (demo)",
        "key_units": ["R-1001 reactor", "C-1001 compressor", "P-1001 pump", "E-1005 heat exchanger",
                      "V-1001 vessel"],
        "data_origin": "synthetic_demo",
    },
}
with open(PLANT, "w", encoding="utf-8") as f:
    json.dump(plant_ctx, f, indent=2, ensure_ascii=False)
print("[ok] plant_context.json written")

# ---------------------------------------------------------------- manifests
MANIFEST = {
    "dataset_version": "2.0.0",
    "generation_timestamp": GEN_ISO,
    "generator": "sovereign synthetic knowledge generator (local, seed=42, no cloud)",
    "source_pid_registry": "data/pid_analysis/registry/verified_equipment_registry.json",
    "anchor_asset": {
        "tag": "R-1001", "equipment_type": "Reactor", "source_drawing": "158.jpg",
        "plant": "Hydrogen Production Plant", "verification_status": "verified",
        "data_origin": "public_pid",
    },
    "documents_generated": [
        {"path": "assets/R-1001/profile.json", "type": "canonical_profile", "data_origin": "synthetic_demo+public_pid"},
        {"path": "assets/R-1001/manual/manual.docx", "type": "equipment_manual", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/sop/operating_sop.docx", "type": "operating_sop", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/sop/pm_sop.docx", "type": "preventive_maintenance_sop", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/inspection/inspection_report.pdf", "type": "inspection_report", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/maintenance/maintenance_history.xlsx", "type": "maintenance_history", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/maintenance/maintenance_history.docx", "type": "maintenance_history_summary", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/sensors/sensor_dataset.csv", "type": "sensor_dataset", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/correspondence/vendor_correspondence.eml", "type": "vendor_correspondence", "data_origin": "synthetic_demo"},
        {"path": "assets/R-1001/approvals/approval_note.docx", "type": "maintenance_approval_note", "data_origin": "synthetic_demo"},
        {"path": "plant/plant_context.json", "type": "plant_context", "data_origin": "synthetic_demo"},
    ],
    "synthetic_data_disclaimer": "ALL operating values, alarms, histories, findings, correspondence "
        "and approvals in this dataset are SYNTHETIC DEMO DATA. They were NOT measured from the "
        "public P&ID and must not be represented as real plant data. Only the asset identity "
        "(R-1001, Reactor, 158.jpg, Hydrogen Production Plant) is taken from the verified P&ID.",
    "relationships": [
        "profile.json is the canon; every document references its synthetic parameters.",
        "sensor_dataset.csv provides the TI-1001/PI-1001/VI-1001 trend that breaches thresholds.",
        "inspection_report.pdf explains the breach (catalyst hotspot + thermowell drift + gasket weep).",
        "operating_sop.docx + pm_sop.docx define the required alarm/trip and corrective actions.",
        "vendor_correspondence.eml recommends the same parts (HRS-CAT-22, HRS-GSK-1001, HRS-TW-1001).",
        "approval_note.docx (APP-2026-0042) summarizes the evidence chain and requests shutdown approval.",
    ],
    "data_origin_policy": "synthetic_demo for generated values; public_pid only for R-1001 identity/context.",
}
with open(os.path.join(ROOT, "metadata", "dataset_manifest.json"), "w", encoding="utf-8") as f:
    json.dump(MANIFEST, f, indent=2, ensure_ascii=False)
print("[ok] dataset_manifest.json written")

GROUND_TRUTH = {
    "asset": "R-1001",
    "scenario": "Gradual, correlated rise in reactor temperature (TI-1001), pressure (PI-1001) and "
                "vibration (VI-1001) over ~35 days, breaching the 320 C high-high temperature alarm, "
                "21 bar pressure-high and 4.0 mm/s vibration-high. Evidence points to catalyst "
                "deactivation (thermal hotspot) plus thermowell drift and a top-head gasket weep. "
                "Inspection confirms; vendor recommends catalyst + gasket replacement and thermowell "
                "recalibration; a controlled shutdown and maintenance approval are required.",
    "expected_findings": [
        "TI-1001 temperature exceeds 310 C (high) by day ~78.7 and 320 C (high-high) by day ~87.6 of the 90-day window (max 322.4 C).",
        "VI-1001 vibration exceeds 4.0 mm/s (high) by day ~75.4 (max 5.3 mm/s).",
        "PI-1001 pressure exceeds 21 bar (high) by day ~77.4 (max 22.4 bar).",
        "Inspection finds catalyst-bed thermal hotspot (activity 61%), thermowell drift ~3.5 C, top-head gasket weep.",
        "Vendor correspondence recommends catalyst HRS-CAT-22 + gasket HRS-GSK-1001 + thermowell HRS-TW-1001.",
    ],
    "expected_action": "Initiate controlled reactor shutdown (ESD per Operating SOP); replace catalyst "
                        "charge and top-head gasket and recalibrate thermowell; do not restart until "
                        "high-high event closed; obtain maintenance approval (shutdown + spares).",
    "evidence_documents": [
        "assets/R-1001/sensors/sensor_dataset.csv",
        "assets/R-1001/inspection/inspection_report.pdf",
        "assets/R-1001/profile.json",
        "assets/R-1001/manual/manual.docx",
        "assets/R-1001/sop/operating_sop.docx",
        "assets/R-1001/sop/pm_sop.docx",
        "assets/R-1001/correspondence/vendor_correspondence.eml",
        "assets/R-1001/approvals/approval_note.docx",
    ],
    "reasoning_required": [
        "Retrieve sensor_dataset.csv and detect multi-signal threshold breach (temperature, pressure, vibration).",
        "Link the breach to the Inspection Report finding (catalyst hotspot + thermowell drift + gasket weep).",
        "Confirm the Operating SOP / PM SOP define the ESD + shutdown + corrective actions.",
        "Confirm the Vendor Correspondence recommends matching parts (HRS-CAT-22, HRS-GSK-1001, HRS-TW-1001).",
        "Summarize the chain in the Approval Note and confirm shutdown/maintenance approval is required (expected_approval=true).",
    ],
    "expected_approval": True,
    "data_origin": "synthetic_demo",
    "note": "Ground-truth for evaluation only. DO NOT ingest into the normal RAG knowledge base.",
}
with open(os.path.join(ROOT, "metadata", "cross_document_ground_truth.json"), "w", encoding="utf-8") as f:
    json.dump(GROUND_TRUTH, f, indent=2, ensure_ascii=False)
print("[ok] cross_document_ground_truth.json written")
print("DONE")
